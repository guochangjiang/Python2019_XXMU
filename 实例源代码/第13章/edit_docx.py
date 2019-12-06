# edit_docx.py
import copy
import docx

doc_file = docx.Document(doc_path)
if 1:
    # 修改段落内容
    for paragraph in doc_file.paragraphs:
        # 深度复制段落内容，包括样式。如果不深度复制，样式会丢失
        list_runs = copy.deepcopy(paragraph.runs)
        paragraph.clear()
        #文字替换
        for run in list_runs:
            for name in dic:
                print (name)                
                if name in run.text:
                    value = dic[name]
                    run.text = run.text.replace(name, str(value))
            # 段落样式的复制
            paragraph.add_run(run.text, run.style)
    # 替换表格内容
    for table in doc_file.tables:
        # 深度复制表格内容，包括样式, 如果不深度复制，样式会丢失
        table_style = copy.deepcopy(table.style)
        for row in table.rows:
            for cell in row.cells:
                for name in dic:
                    if name in cell.text:
                        value = dic[name]
                        cell.text = cell.text.replace(name, str(value))
        table.style=table_style
    doc_file.save(new_doc_path)
    print(new_doc_path)
    # 删除段落内容
    for line in range(len(doc_file.paragraphs)):
        if line in delete_paragraphs:
            doc_file.paragraphs[line].clear()
        for run in range(doc_file.paragraphs[line].runs):
            if run in delete_runs:
                doc_file.paragraphs[line].runs[run].clear()
    # 删除表格内容
    for table in range(len(doc_file.tables)):
        if table in delete_tables:
            doc_file.tables[table].clear()
        for row in range(doc_file.tables[table].runs):
            if row in delete_rows:
                doc_file.tables[table].rows[row].clear()
            for cell in range(doc_file.tables[table].rows[row].cells):
                if cell in delete_cells:
                    doc_file.tables[table].rows[row].cells[cell].clear()
    doc_file.save(new_doc_path)
    # 增加段落内容
    doc_file.add_paragraph("段落文字").add_run("run文字")
    doc_file.paragraphs[1].add_run("run文字")
    # 增加表格内容
    for table in range(len(doc_file.tables)):
        doc_file.add_table(2,3)
    doc_file.tables[1].add_row('文字')