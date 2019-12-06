# merge_docx.py
import docx

docx_list = ("CV_1.docx", "CV_2.docx", "CV_3.docx")
# 新建一个空的word文档对象
doc = docx.Document()
for docx_file in docx_list:
    for para in docx.Document(docx_file).paragraphs:
        doc.add_paragraph(para.text)
# 保存新的word文档
doc.save("CV_merge1.docx")

merged_document = docx.Document()
for index, file in enumerate(docx_list):
    sub_doc = docx.Document(file)
    # Don't add a page break if you've reached the last file.
    if index < len(docx_list)-1:
        sub_doc.add_page_break()
    for element in sub_doc.element.body:
        merged_document.element.body.append(element)
merged_document.save('CV_merge2.docx')