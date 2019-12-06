# create_docx.py
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()# 设置一个空白样式
style = document.styles['Normal']
# 设置西文字体
style.font.name = 'Times New Roman'
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置标题
title_ = document.add_heading(level=0)
# 标题居中
title_.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 添加标题内容
title_run = title_.add_run("文档标题（title）")
# 设置标题字体大小
title_run.font.size = Pt(20)
# 设置标题西文字体
title_run.font.name = 'Times New Roman'
# 设置标题中文字体
title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
# document.add_heading('文档标题', 0)

# 获取段落样式
paragraph_format = style.paragraph_format
# 首行缩进0.74厘米，即2个字符
paragraph_format.first_line_indent = Cm(0.74)

p = document.add_paragraph('本段落的文字有')
p.add_run('加粗').bold = True
p.add_run('和')
p.add_run('斜体。').italic = True

hl1=document.add_heading(level=1).add_run('一级标题')
hl1.font.name = 'Times New Roman'
hl1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
document.add_heading('二级标题', level=2)
document.add_heading('三级标题', level=3)
document.add_paragraph('引用', style='Intense Quote')
document.add_paragraph('无序列表1', style='List Bullet')
document.add_paragraph('无序列表2', style='List Bullet')
document.add_paragraph('有序列表1', style='List Number')
document.add_paragraph('有序列表2', style='List Number')

picture = document.add_paragraph()
picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
picture.add_run().add_picture('Python.png', width=Inches(3))

records = (
    (1, '第一', 'Python'),
    (2, '第二', 'Java'),
    (3, '其他', 'C, R, Go')
)
table = document.add_table(rows=1, cols=3, style='Medium Grid 1 Accent 1')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '序号'
hdr_cells[1].text = '排名'
hdr_cells[2].text = '语言'
for order, rank, lang in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(order)
    row_cells[1].text = rank
    row_cells[2].text = lang

document.add_page_break()
document.save('demo.docx')