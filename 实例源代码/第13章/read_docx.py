# read_docs.py

import docx
from docx import Document

path = "test.docx"
document = Document(path)
# 按段落读取文字
for paragraph in document.paragraphs:
    print(paragraph.text)
# 读取run
for i in range(0, len(document.paragraphs[1].runs)):
    print(document.paragraphs[1].runs[i].text)
# 读取表格
for table in document.tables:
    for i in range(0, len(table.rows)):
        print("{} {} {}".format(table.cell(i,0).text, table.cell(i,1).text, table.cell(i,2).text))